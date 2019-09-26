# -*- coding: utf-8 -*-
from bs4 import BeautifulSoup
import xlwings as xw
import os
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.shared import Pt

## 兼容问题，修改路径兼容mac
#url = os.getcwd() + '/'
url = os.getcwd() + '\\'

vus = []
gaoshu = []
zhongshu = []
dishu = []
zongshu = []

soup = BeautifulSoup(open('index.html', encoding='utf8'),features='html.parser')
hosts = soup.select('#content>div:nth-child(6)>div:nth-child(2)>table>tbody>tr>td>a')
print (hosts)
for host in hosts:
    hos = host.get('href')
    hhh = host.get_text()
    print (host)
    print (hhh)
    print (type(host))


    soup = BeautifulSoup(open(hos ,encoding='utf8'),features="html.parser")
    vuls = soup.select('#vuln_list > tbody > tr > td > ul > li > div > span')

    val = []
    low_list = []
    middle_list = []
    high_list = []
    for vul in vuls:
        # 威胁等级
        level = vul.get('class')[0]
        # 漏洞名称
        v = vul.get_text()
        # 去重
        if v in val:
            continue
        # 漏洞详情
        val.append(v)

        ## 漏洞等级分类并替换漏洞等级为中文
        if level == 'level_danger_low':
            low_list.append( {
            'level_danger' : '低',
            'vul' : v,
        })
        elif level == 'level_danger_middle':
            middle_list.append( {
            'level_danger' : '中',
            'vul' : v,
        })
        else:
            high_list.append( {
            'level_danger' : '高',
            'vul' : v,
        })

        # 漏洞个数
        gao = len(high_list)
        gaoshu.append(gao)
        zhong = len(middle_list)
        #zhongshu.append(zhong)
        di = len(low_list)
        #dishu.append(di)
        zong = len(high_list) + len(middle_list) + len(low_list)
        #zongshu.append(zong)
        zz = high_list + middle_list + low_list


    print (zz)

    print (type(zz))

    print (gao)
    print (zhong)
    print (di)
    print (zong)


    for z in zz:
        print (z['vul'])
        print (z['level_danger'])


    # 添加标题
    # dd = document.add_heading(hhh + '漏洞扫描报告', 0)
    # dd.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 添加第一段落
    # document.add_heading('一、漏洞统计', 1)

    document = Document()
    #设置全局样式
    document.styles['Normal'].font.name=u'黑体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')

    # 首标题
    p_h = document.add_paragraph()
    run = p_h.add_run(hhh + '漏洞扫描报告')
    run.font.size = Pt(26)
    p_h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER 
    run.bold = True
    run.font.color.rgb = RGBColor(0,0,0)

    # 添加分页
    document.add_page_break()

    # 标题一
    p_1 = document.add_paragraph()
    run = p_1.add_run('一、漏洞统计')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0,0,0)

    table = document.add_table(rows=2, cols=4, style="Light List")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '高危'
    hdr_cells[1].text = '中危'
    hdr_cells[2].text = '低危'
    hdr_cells[3].text = '合计'

    hdc_cells = table.rows[1].cells
    hdc_cells[0].text = str(gao)
    hdc_cells[1].text = str(zhong)
    hdc_cells[2].text = str(di)
    hdc_cells[3].text = str(zong)

    #添加第二段落
    p_2 = document.add_paragraph()
    run = p_2.add_run('二、漏洞详情')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0,0,0)

    table = document.add_table(rows=zong+1, cols=3, style="Light List")
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '危险程度'
    hdr_cells[2].text = '漏洞名称'

    for z, element in enumerate(zz):
        hdr_cells = table.rows[z+1].cells
        hdr_cells[0].text = str(z+1)
        hdr_cells[1].text = element['level_danger']
        hdr_cells[2].text = element['vul']

    document.save(hhh +'.docx')
print('\n---------------------------------------------------------------------------\n')
